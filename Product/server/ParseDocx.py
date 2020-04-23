import docx
import os

def ParseDocx(docxPath):
    f = open(docxPath, 'rb')
    document = docx.Document(f)
    f.close()

    txt_file_names = []
    txt_file_contents = []

    # Retrieve JURISDICTION Section
    start_idx = next((i for i,v in enumerate(document.paragraphs) if v.text == "JURISDICTION"), None)
    if start_idx != None:
        start_idx += 1
        end_idx = start_idx + next((i for i,v in enumerate(document.paragraphs[start_idx:]) if v.text.isupper()))
        paragraph_text = ' '.join([x.text for x in document.paragraphs[start_idx:end_idx]])
        txt_file_names.append("JURISDICTION.txt")
        txt_file_contents.append(paragraph_text)

    # Retrieve COUNTERSTATEMENT OF THE ISSUES Section
    start_idx = next((i for i,v in enumerate(document.paragraphs) if v.text == "COUNTERSTATEMENT OF THE ISSUES"), None)
    if start_idx != None:
        start_idx += 1
        end_idx = start_idx + next((i for i,v in enumerate(document.paragraphs[start_idx:]) if v.text.isupper()))
        paragraph_text = ' '.join([x.text for x in document.paragraphs[start_idx:end_idx]])
        txt_file_names.append("COUNTERSTATEMENT OF THE ISSUES.txt")
        txt_file_contents.append(paragraph_text)

    # Retrieve STATEMENT OF THE CASE Section
    start_idx = next((i for i,v in enumerate(document.paragraphs) if v.text == "STATEMENT OF THE CASE"), None)
    if start_idx != None:
        start_idx += 1
        end_idx = start_idx + next((i for i,v in enumerate(document.paragraphs[start_idx:]) if v.text.isupper()))
        paragraph_text = ' '.join([x.text for x in document.paragraphs[start_idx:end_idx]])
        txt_file_names.append("STATEMENT OF THE CASE.txt")
        txt_file_contents.append(paragraph_text)

    # Retrieve STATEMENT OF ISSUES Section
    start_idx = next((i for i,v in enumerate(document.paragraphs) if v.text == "STATEMENT OF ISSUES"), None)
    if start_idx != None:
        start_idx += 1
        end_idx = start_idx + next((i for i,v in enumerate(document.paragraphs[start_idx:]) if v.text.isupper()))
        paragraph_text = ' '.join([x.text for x in document.paragraphs[start_idx:end_idx]])
        txt_file_names.append("STATEMENT OF ISSUES.txt")
        txt_file_contents.append(paragraph_text)

    # Retrieve SUMMARY OF ARGUMENT Section
    start_idx = next((i for i,v in enumerate(document.paragraphs) if v.text == "SUMMARY OF ARGUMENT"), None)
    if start_idx != None:
        start_idx += 1
        end_idx = start_idx + next((i for i,v in enumerate(document.paragraphs[start_idx:]) if v.text.isupper()))
        paragraph_text = ' '.join([x.text for x in document.paragraphs[start_idx:end_idx]])
        txt_file_names.append("SUMMARY OF ARGUMENT.txt")
        txt_file_contents.append(paragraph_text)

    # Retrieve STANDARD OF REVIEW Section
    start_idx = next((i for i,v in enumerate(document.paragraphs) if v.text == "STANDARD OF REVIEW"), None)
    if start_idx != None:
        start_idx += 1
        end_idx = start_idx + next((i for i,v in enumerate(document.paragraphs[start_idx:]) if v.text.isupper()))
        paragraph_text = ' '.join([x.text for x in document.paragraphs[start_idx:end_idx]])
        txt_file_names.append("STANDARD OF REVEIW.txt")
        txt_file_contents.append(paragraph_text)

    # Retrieve ARGUMENT Sections
    arg_number = 1
    start_idx = next((i for i,v in enumerate(document.paragraphs[end_idx:]) if v.text == "ARGUMENT"), None)
    if start_idx != None:
        start_idx = start_idx + end_idx
        while start_idx != None:
            roman_num = int2roman(arg_number)
            next_start_idx = next((i for i,v in enumerate(document.paragraphs[start_idx:]) if v.text.startswith('{}.'.format(roman_num))), None)
            if next_start_idx != None:
                start_idx += next_start_idx
                if start_idx != None:
                    start_idx += 1
                    end_idx = next((i for i,v in enumerate(document.paragraphs[start_idx:]) if v.text.startswith('{}.'.format(int2roman(arg_number + 1)))), None)
                    if end_idx == None:
                        end_idx = start_idx + next((i for i, v in enumerate(document.paragraphs[start_idx:]) if v.text.isupper()))
                    else:
                        end_idx += start_idx
                    if end_idx != None:
                        paragraph_text = ' '.join([x.text for x in document.paragraphs[start_idx:end_idx]])
                        txt_file_names.append("ARGUMENT {}.txt".format(roman_num))
                        txt_file_contents.append(paragraph_text)
                        arg_number += 1
                        start_idx = end_idx
                    else:
                        start_idx = None
            else:
                start_idx = None

    dir = os.path.dirname(docxPath)
    new_folder = os.path.splitext(os.path.basename(docxPath))[0]
    if not os.path.exists(dir + "/" + new_folder):
        os.makedirs(dir + "/" + new_folder)
    for filename, filecontents in zip(txt_file_names, txt_file_contents):
        f = open(dir + "/" + new_folder + "/" + filename, "w")
        f.write(filecontents)
        f.close()

def int2roman(number):
    """Code retrieved from: https://www.daniweb.com/programming/software-development/code/216865/roman-numerals-python"""
    numerals = { 1 : "I", 4 : "IV", 5 : "V", 9 : "IX", 10 : "X", 40 : "XL",
        50 : "L", 90 : "XC", 100 : "C", 400 : "CD", 500 : "D", 900 : "CM", 1000 : "M" }
    result = ""
    for value, numeral in sorted(numerals.items(), reverse=True):
        while number >= value:
            result += numeral
            number -= value
    return result

def main():
    ParseDocx("/Users/brandon/Desktop/Parsing/Apple Answering Brief - Word Version.docx")

if __name__ == "__main__":
    main()
